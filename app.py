import os
import re
import uuid
import traceback
from pathlib import Path
from datetime import datetime
from flask import Flask, request, jsonify, send_file, send_from_directory
from flask_cors import CORS
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

app = Flask(__name__, static_folder=".", static_url_path="")
CORS(app)

UPLOAD_FOLDER = "uploads_temp"
OUTPUT_FOLDER = "outputs"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc", ".xlsx", ".xls", ".xlsm", ".png", ".jpg", ".jpeg", ".tiff", ".bmp"}

# ── Date normalization ─────────────────────────────────
MONTH_MAP = {
    'jan':'01','feb':'02','mar':'03','apr':'04','may':'05','jun':'06',
    'jul':'07','aug':'08','sep':'09','oct':'10','nov':'11','dec':'12'
}

def normalize_date(raw):
    if not raw: return ''
    raw = raw.strip()
    raw = re.sub(r'(\d+)(st|nd|rd|th)', r'\1', raw, flags=re.IGNORECASE)
    formats = [
        "%d/%m/%Y","%d-%m-%Y","%d.%m.%Y",
        "%m/%d/%Y","%m-%d-%Y",
        "%d/%m/%y","%d-%m-%y",
        "%Y-%m-%d","%Y/%m/%d",
        "%d %B %Y","%d %b %Y",
        "%B %d, %Y","%b %d, %Y",
        "%B %d %Y","%b %d %Y",
    ]
    for fmt in formats:
        try: return datetime.strptime(raw, fmt).strftime("%Y-%m-%d")
        except ValueError: pass
    m = re.match(r"(\d{1,2})\s+([A-Za-z]+)\.?\s+(\d{4})", raw)
    if m:
        mon = MONTH_MAP.get(m.group(2)[:3].lower())
        if mon:
            try: return datetime.strptime(f"{m.group(1)}/{mon}/{m.group(3)}", "%d/%m/%Y").strftime("%Y-%m-%d")
            except: pass
    return raw

# ── Regex patterns ─────────────────────────────────────

GSTIN_RE = re.compile(r'\b\d{2}[A-Z]{5}\d{4}[A-Z][1-9A-Z]Z[0-9A-Z]\b')

def strip_gstin(text):
    """Remove GSTIN numbers from text before invoice no extraction to avoid false matches."""
    return GSTIN_RE.sub('', text)

# ── Invoice Number ──
INVOICE_NO_PATTERNS = [
    # "Invoice No. INV20240001" on same line
    r"invoice\s*(?:no|number)?\.?\s*[:\-]?\s*([A-Z]{1,5}\d{5,})",
    r"invoice\s*(?:no|number)?\.?\s*[:\-]?\s*(\d{8,})",
    # Same line with lookahead
    r"invoice\s*no\.?\s*[:\-]?\s*(\d+|[A-Z0-9][A-Z0-9\-\/\_\.]*)(?=\s+(?:invoice|bill|due|date|total|tax|ship|from|to)\b|\s*$|\s*\n)",
    r"invoice\s*number\s*[:\-]?\s*(\d+|[A-Z0-9][A-Z0-9\-\/\_\.]*)(?=\s+(?:invoice|bill|due|date|total|tax|ship|from|to)\b|\s*$|\s*\n)",
    r"invoice\s*#\s*[:\-]?\s*(\d+|[A-Z0-9][A-Z0-9\-\/\_\.]*)(?=\s+[A-Za-z]|\s*$|\s*\n)",
    r"inv\s*no\.?\s*[:\-]?\s*(\d+|[A-Z0-9][A-Z0-9\-\/\_\.]*)(?=\s+[A-Za-z]|\s*$|\s*\n)",
    r"bill\s*no\.?\s*[:\-]?\s*(\d+|[A-Z0-9][A-Z0-9\-\/\_\.]*)(?=\s+[A-Za-z]|\s*$|\s*\n)",
    r"bill\s*number\s*[:\-]?\s*(\d+|[A-Z0-9][A-Z0-9\-\/\_\.]*)(?=\s+[A-Za-z]|\s*$|\s*\n)",
    r"tax\s*invoice\s*(?:no\.?|number|#)?\s*[:\-]?\s*(\d+|[A-Z0-9][A-Z0-9\-\/\_\.]*)(?=\s+[A-Za-z]|\s*$|\s*\n)",
    r"receipt\s*(?:no\.?|number|#)\s*[:\-]?\s*(\d+|[A-Z0-9][A-Z0-9\-\/\_\.]*)(?=\s+[A-Za-z]|\s*$|\s*\n)",
    # Next line: "Invoice No.\nINV20240001"
    r"invoice\s*no\.?\s*[:\-]?\s*\n\s*([A-Z0-9][A-Z0-9\-\/\_\.]*)",
    r"invoice\s*number\s*[:\-]?\s*\n\s*([A-Z0-9][A-Z0-9\-\/\_\.]*)",
    r"inv\.?\s*no\.?\s*[:\-]?\s*\n\s*([A-Z0-9][A-Z0-9\-\/\_\.]*)",
    r"bill\s*no\.?\s*[:\-]?\s*\n\s*([A-Z0-9][A-Z0-9\-\/\_\.]*)",
    # Named prefixes (standalone)
    r"\b(INV[-\/]?\d{4,}[A-Z0-9]*)\b",
    r"\b(RD\d{5,})\b",
    r"\b(BILL[-\/]?[0-9A-Z\-\_]{1,20})\b",
]

# ── Invoice Date ──
INVOICE_DATE_PATTERNS = [
    # Same line
    r"invoice\s*date\s*[:\-]?\s*(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})",
    r"inv\.?\s*date\s*[:\-]?\s*(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})",
    r"bill\s*date\s*[:\-]?\s*(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})",
    r"date\s*of\s*(?:invoice|bill)\s*[:\-]?\s*(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})",
    r"dated?\s*[:\-]?\s*(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})",
    r"date\s*[:\-]?\s*(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})",
    # Next line
    r"invoice\s*date\s*[:\-]?\s*\n\s*(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})",
    r"date\s*[:\-]?\s*\n\s*(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})",
    # ISO
    r"(\d{4}[-\/]\d{2}[-\/]\d{2})",
    # Text months
    r"(\d{1,2}(?:st|nd|rd|th)?\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4})",
    r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{1,2}(?:st|nd|rd|th)?,?\s+\d{4})",
]

# ── Quantity / kWh / Units ──
QUANTITY_PATTERNS = [
    # "2,580.0 kWh" preceded by "Energy Charges"
    r"energy\s*charges?\s+(\d[\d,]*\.?\d*)\s*kwh",
    # "Units Consumed (kWh)" in table header, value is on same/next row as big number
    # Match "2,580.0" as standalone kWh figure (4+ digits with optional comma/decimal)
    r"([\d,]{4,}\.\d+)\s*kwh",
    r"(\d[\d,]+\.\d)\s*kwh",
    # "consumption: 2580"
    r"(?:total\s*)?consumption\s*[:\-]?\s*([\d,]+\.?\d*)\s*(?:kwh|kw|units)",
    # "units consumed: 2580"
    r"units?\s*consumed\s*[:\-]?\s*([\d,]+\.\d+)",
    # Generic qty/quantity
    r"(?:qty|quantity)\s*[:\-]?\s*(\d[\d,]*\.?\d*)",
]

# ── Cost / Total Amount ──
COST_PATTERNS = [
    # Grand total AED 1,155.42
    r"grand\s*total\s*(?:AED|INR|USD|EUR|GBP|Rs\.?|₹|\$|£|€|R)?\s*([\d,]+\.\d{2})",
    # "AED 1,155.42" — explicit AED currency
    r"(?:AED|INR|USD|EUR|GBP)\s*([\d,]+\.\d{2})\b",
    # AMOUNT DUE + newlines + AED value
    r"amount\s*due[\s\S]{0,80}?(?:AED|INR|USD|EUR|GBP|Rs\.?|₹|\$|£|€)\s*([\d,]+\.\d{2})",
    # Grand total / total amount / amount due with optional currency
    r"(?:grand\s*total|total\s*amount|amount\s*due|net\s*amount|total\s*payable|balance\s*due|total\s*cost|net\s*bill\s*amount)\s*[:\-]?\s*(?:AED|INR|USD|EUR|GBP|Rs\.?|₹|\$|£|€|R)?\s*([\d,]+\.?\d*)",
    # Total / subtotal
    r"(?:total|subtotal|sub\s*total)\s*[:\-]?\s*(?:AED|INR|USD|EUR|GBP|Rs\.?|₹|\$|£|€|R)?\s*([\d,]+\.?\d*)",
    # ₹ / $ / £ / € prefix
    r"[₹\$£€]\s*([\d,]+\.?\d{0,2})\b",
    # Rs. 25.00
    r"(?:rs\.?|inr|usd|eur|gbp|aed)\s*([\d,]+\.?\d*)",
    # Fallback: amount followed by total/due/payable
    r"([\d,]+\.\d{2})\s*(?:total|due|payable|only)",
]

def first_match(patterns, text):
    for p in patterns:
        m = re.search(p, text, re.IGNORECASE | re.MULTILINE)
        if m: return m.group(1).strip()
    return ""

def extract_header_value_row(text):
    """Handle the tabular header layout:
    Line 1: 'Invoice No. Invoice Date Due Date Account No.'
    Line 2: 'INV20240001 05/01/2024  26/01/2024  AC-00112233'
    Returns (invoice_no, invoice_date) or ("", "").
    """
    m = re.search(
        r'Invoice\s+No\.?\s+Invoice\s+Date.*?\n'
        r'\s*([A-Z0-9][A-Z0-9\-\/]*?)\s+'
        r'(\d{1,2}[\/-]\d{1,2}[\/-]\d{2,4})',
        text, re.IGNORECASE
    )
    if m:
        return m.group(1).strip(), m.group(2).strip()
    return "", ""

def calculate_confidence(fields):
    """Calculate a confidence percentage based on which fields were extracted."""
    score = 0
    if fields.get("Invoice No."):  score += 25
    if fields.get("Invoice Date"): score += 25
    if fields.get("Cost"):         score += 25
    if fields.get("Quantity"):     score += 15
    # Heuristics for quality
    inv = fields.get("Invoice No.", "")
    if inv and re.match(r'^[A-Z]{2,5}\d{5,}$', inv): score += 5
    if fields.get("Invoice Date") and "-" in fields["Invoice Date"]: score += 5
    return min(score, 100)

def extract_fields(text):
    """Extract invoice fields from a block of text (one invoice)."""
    clean = strip_gstin(text)

    # ── Try header-value row approach first (tabular PDFs) ──
    hdr_inv, hdr_date = extract_header_value_row(text)

    # ── Regex-based extraction ──
    inv_no       = hdr_inv or first_match(INVOICE_NO_PATTERNS, clean)
    inv_date_raw = hdr_date or first_match(INVOICE_DATE_PATTERNS, text)
    inv_date     = normalize_date(inv_date_raw)
    qty          = first_match(QUANTITY_PATTERNS, text)
    cost         = first_match(COST_PATTERNS, text)

    # ── Fallback: INV prefix anywhere ──
    if not inv_no:
        m = re.search(r'\bINV[-/]?\d{4,}[A-Z0-9]*\b', clean, re.IGNORECASE)
        if m: inv_no = m.group(0)

    # ── Sanitize bad matches ──
    if inv_no and inv_no.lower() in ["invoice", "number", "no", "date"]: inv_no = ""

    # ── De-duplicate: if qty == cost, clear qty (it grabbed the total) ──
    if qty and cost and qty.replace(",","") == cost.replace(",",""):
        qty = ""

    fields = {
        "Invoice No.":  inv_no,
        "Invoice Date": inv_date,
        "Quantity":     qty,
        "Cost":         cost,
    }
    fields["Confidence Score"] = f"{calculate_confidence(fields)}%"
    return fields

# ── Text extraction per file type ──────────────────────
def get_text_chunks(filepath):
    """Returns a list of text blocks — one per page for PDFs, one total for others."""
    ext = Path(filepath).suffix.lower()
    chunks = []
    if ext == ".pdf":
        import pdfplumber
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        row_text = " | ".join(str(c) for c in row if c)
                        if row_text.strip(): text += "\n" + row_text
                # OCR fallback for scanned pages
                if len(text.strip()) < 50:
                    try:
                        from pdf2image import convert_from_path
                        import pytesseract
                        images = convert_from_path(filepath, first_page=page.page_number, last_page=page.page_number, dpi=200)
                        for img in images:
                            text += "\n" + pytesseract.image_to_string(img, lang="eng")
                    except: pass
                if text.strip():
                    chunks.append(text)
    elif ext in (".docx", ".doc"):
        from docx import Document
        doc = Document(filepath)
        lines = [p.text for p in doc.paragraphs]
        for tbl in doc.tables:
            for row in tbl.rows:
                lines.append(" | ".join(c.text.strip() for c in row.cells))
        chunks.append("\n".join(lines))
    elif ext in (".xlsx", ".xls", ".xlsm"):
        wb = openpyxl.load_workbook(filepath, data_only=True)
        lines = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                r = " | ".join(str(c) for c in row if c is not None)
                if r.strip(): lines.append(r)
        chunks.append("\n".join(lines))
    elif ext in (".png", ".jpg", ".jpeg", ".tiff", ".bmp"):
        import pytesseract
        from PIL import Image
        chunks.append(pytesseract.image_to_string(Image.open(filepath), lang="eng"))
    else:
        raise ValueError(f"Unsupported file type: {ext}")
    return chunks

# ── Template Excel export (Envizi format) ──────────────
TEMPLATE_COLS = [
    "Organization Link", "Organization", "Location", "Location Ref",
    "Account Style Link", "Account Style Caption", "Account Subtype",
    "Account Number", "Account Reference", "Account Supplier", "Account Reader",
    "Record Start YYYY-MM-DD", "Record End YYYY-MM-DD", "Record Data Quality",
    "Record Billing Type", "Record Subtype", "Record Entry Method",
    "Record Reference", "Record Invoice Number", "Quantity", "Total Cost"
]

TEMPLATE_FIXED = {
    "Organization Link":     17000150,
    "Organization":          "DHI.AI Demo",
    "Location":              "UAE Manufacturing",
    "Location Ref":          "",
    "Account Style Link":    14445,
    "Account Style Caption": "S2 - Electricity - kWh",
    "Account Subtype":       "",
    "Account Number":        "Electricity Consumption UAE",
    "Account Reference":     "",
    "Account Supplier":      "",
    "Account Reader":        "",
    "Record Data Quality":   "",
    "Record Billing Type":   "",
    "Record Subtype":        "",
    "Record Entry Method":   "",
    "Record Reference":      "",
}

def build_template_excel(records, output_path):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Account_Setup_and_Data_Load_-_PM&C"

    hdr_fill = PatternFill("solid", fgColor="1E3A5F")
    hdr_font = Font(bold=True, color="FFFFFF", name="Calibri", size=10)
    center   = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin     = Side(style="thin", color="AAAAAA")
    border   = Border(left=thin, right=thin, top=thin, bottom=thin)

    col_widths = {
        "Organization Link": 18, "Organization": 16, "Location": 20,
        "Account Style Link": 18, "Account Style Caption": 30,
        "Account Number": 26, "Account Reference": 28,
        "Record Start YYYY-MM-DD": 22, "Record End YYYY-MM-DD": 20,
        "Record Invoice Number": 22, "Quantity": 14, "Total Cost": 14,
    }

    for ci, col in enumerate(TEMPLATE_COLS, 1):
        cell = ws.cell(row=1, column=ci, value=col)
        cell.font, cell.fill, cell.alignment, cell.border = hdr_font, hdr_fill, center, border
        ws.column_dimensions[openpyxl.utils.get_column_letter(ci)].width = col_widths.get(col, 16)
    ws.row_dimensions[1].height = 36

    alt_fill = PatternFill("solid", fgColor="F0F4FA")
    for ri, rec in enumerate(records, 2):
        fill = alt_fill if ri % 2 == 0 else PatternFill("solid", fgColor="FFFFFF")
        row_data = dict(TEMPLATE_FIXED)
        row_data["Record Start YYYY-MM-DD"] = rec.get("Invoice Date", "")
        row_data["Record End YYYY-MM-DD"]   = rec.get("Invoice Date", "")
        row_data["Record Invoice Number"]   = rec.get("Invoice No.", "")

        # Quantity → float
        raw_qty = rec.get("Quantity", "")
        try:    row_data["Quantity"] = float(str(raw_qty).replace(",", "").strip()) if raw_qty else ""
        except: row_data["Quantity"] = raw_qty

        # Total Cost → float
        raw_cost = rec.get("Cost", "")
        try:    row_data["Total Cost"] = float(str(raw_cost).replace(",", "").strip()) if raw_cost else ""
        except: row_data["Total Cost"] = raw_cost

        for ci, col in enumerate(TEMPLATE_COLS, 1):
            cell = ws.cell(row=ri, column=ci, value=row_data.get(col, ""))
            cell.font   = Font(name="Calibri", size=10)
            cell.fill   = fill
            cell.border = border
            cell.alignment = Alignment(vertical="center")
            if col in ("Total Cost", "Quantity") and isinstance(cell.value, float):
                cell.number_format = '#,##0.00'

    ws.freeze_panes = "A2"
    wb.save(output_path)

# ── Preview Excel (selected columns) ──────────────────
def build_preview_excel(records, selected_cols, output_path):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Extracted Invoices"
    all_cols = ["Invoice No.", "Invoice Date", "Quantity", "Cost"]
    cols = [c for c in all_cols if c in selected_cols]

    hdr_fill = PatternFill("solid", fgColor="1E3A5F")
    hdr_font = Font(bold=True, color="FFFFFF", name="Calibri", size=10)
    border   = Border(
        left=Side(style="thin", color="AAAAAA"), right=Side(style="thin", color="AAAAAA"),
        top=Side(style="thin", color="AAAAAA"),  bottom=Side(style="thin", color="AAAAAA")
    )
    widths = {"Invoice No.": 24, "Invoice Date": 16, "Quantity": 16, "Cost": 18}

    for ci, col in enumerate(cols, 1):
        cell = ws.cell(row=1, column=ci, value=col)
        cell.font, cell.fill, cell.border = hdr_font, hdr_fill, border
        ws.column_dimensions[openpyxl.utils.get_column_letter(ci)].width = widths.get(col, 20)

    for ri, rec in enumerate(records, 2):
        for ci, col in enumerate(cols, 1):
            cell = ws.cell(row=ri, column=ci, value=rec.get(col, ""))
            cell.border = border

    ws.freeze_panes = "A2"
    wb.save(output_path)

# ── Routes ─────────────────────────────────────────────
@app.route("/")
def index():
    return send_from_directory(".", "invoice_ui.html")

@app.route("/extract", methods=["POST"])
def extract():
    uploaded = request.files.getlist("files")
    selected_cols = request.form.getlist("columns")
    if not selected_cols: selected_cols = ["Invoice No.", "Invoice Date", "Quantity", "Cost"]

    if not uploaded: return jsonify({"error": "No files uploaded"}), 400

    records = []
    errors  = []

    for f in uploaded:
        ext = Path(f.filename).suffix.lower()
        if ext not in ALLOWED_EXTENSIONS:
            errors.append({"filename": f.filename, "error": "Unsupported file type"})
            continue

        tmp_path = os.path.join(UPLOAD_FOLDER, f"{uuid.uuid4().hex}{ext}")
        f.save(tmp_path)

        try:
            chunks = get_text_chunks(tmp_path)
            for idx, text in enumerate(chunks):
                if not text.strip(): continue
                fields = extract_fields(text)
                fields["filename"] = f.filename if len(chunks) == 1 else f"{f.filename} (pg {idx+1})"
                records.append(fields)

                print(f"\n-- [{fields['filename']}] --")
                print(f"   Invoice No.: {fields['Invoice No.']}")
                print(f"   Date:        {fields['Invoice Date']}")
                print(f"   Quantity:    {fields['Quantity']}")
                print(f"   Cost:        {fields['Cost']}")
                print(f"   Confidence:  {fields['Confidence Score']}")
        except Exception as e:
            traceback.print_exc()
            errors.append({"filename": f.filename, "error": str(e)})
        finally:
            try: os.remove(tmp_path)
            except: pass

    # Build preview Excel (selected cols)
    out_name = f"extracted_{uuid.uuid4().hex[:8]}.xlsx"
    build_preview_excel(records, selected_cols, os.path.join(OUTPUT_FOLDER, out_name))

    # Build template Excel (Envizi 21-column format)
    tpl_name = f"template_{uuid.uuid4().hex[:8]}.xlsx"
    build_template_excel(records, os.path.join(OUTPUT_FOLDER, tpl_name))

    return jsonify({
        "records":     records,
        "errors":      errors,
        "excel_id":    out_name,
        "template_id": tpl_name,
        "total":       len(records),
    })

@app.route("/download/<filename>")
def download(filename):
    path = os.path.join(OUTPUT_FOLDER, filename)
    if not os.path.exists(path): return jsonify({"error": "File not found"}), 404
    dl_name = "Account_Setup_and_Data_Load_-_PM&C-Elec.xlsx" if filename.startswith("template_") else "extracted_invoices.xlsx"
    return send_file(path, as_attachment=True, download_name=dl_name)

# if __name__ == "__main__":
#     print("\n[*] InvoiceLens running at: http://localhost:5000\n")
#     app.run(debug=True, port=5000)
if __name__ == "__main__":
    import os
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)